import sys
import os

# Install python-docx if not present
try:
    import docx
except ImportError:
    import subprocess
    print("Installing python-docx library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # Page Margins (A4, 1 inch all around)
    for section in doc.sections:
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69) # A4 height
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Charcoal/slate color
    
    # Helper to set font on run
    def set_font(run, font_name='Arial', size_pt=11, bold=False, italic=False, color_rgb=None):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb

    # 1. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run("JUDUL ARTIKEL DITULIS DENGAN HURUF ARIAL BOLD 14 PT\nMAKSIMAL 15 KATA")
    set_font(run_title, size_pt=14, bold=True, color_rgb=RGBColor(0x1A, 0x20, 0x2C))

    # 2. Author Block
    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_authors.paragraph_format.space_after = Pt(4)
    run_auths = p_authors.add_run("Penulis Pertama1*, Penulis Kedua2, Penulis Ketiga3")
    set_font(run_auths, size_pt=10, bold=True, color_rgb=RGBColor(0x2D, 0x37, 0x48))
    
    # Author Affiliations
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(12)
    run_aff = p_aff.add_run("1Jurusan Manajemen, Universitas Negeri Jakarta, Indonesia\n"
                           "2Departemen Ilmu Hukum, Universitas Indonesia, Indonesia\n"
                           "3Fakultas Pertanian, Universitas Gadjah Mada, Indonesia\n"
                           "*Email Korespondensi: penulis_pertama@email.com")
    set_font(run_aff, size_pt=9, italic=True)

    # 3. Abstract Section (Bahasa Indonesia)
    p_abs_title_id = doc.add_paragraph()
    p_abs_title_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs_title_id.paragraph_format.space_before = Pt(12)
    p_abs_title_id.paragraph_format.space_after = Pt(6)
    run_abs_title_id = p_abs_title_id.add_run("ABSTRAK")
    set_font(run_abs_title_id, size_pt=10, bold=True, color_rgb=RGBColor(0x1A, 0x20, 0x2C))

    p_abs_id = doc.add_paragraph()
    p_abs_id.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs_id.paragraph_format.left_indent = Inches(0.5)
    p_abs_id.paragraph_format.right_indent = Inches(0.5)
    p_abs_id.paragraph_format.space_after = Pt(6)
    p_abs_id.paragraph_format.line_spacing = 1.0
    run_abs_id = p_abs_id.add_run(
        "Abstrak ditulis menggunakan bahasa Indonesia yang baik dan benar dengan jenis huruf Arial 10 pt. "
        "Abstrak ditulis secara ringkas dan padat dalam satu paragraf tunggal (tanpa indentasi baris pertama), "
        "berkisar antara 150 sampai 250 kata. Abstrak harus memuat secara lengkap ringkasan penelitian yang terdiri dari: "
        "(1) Latar belakang masalah dan tujuan utama penelitian; (2) Metode penelitian yang mencakup desain, subjek/sampel, "
        "dan instrumen pengumpulan data; (3) Hasil temuan ilmiah utama beserta analisis data; dan (4) Kesimpulan akhir serta "
        "implikasi penting dari hasil temuan tersebut. Hindari menuliskan singkatan, rumus matematika, kutipan pustaka, "
        "atau rujukan ke tabel/gambar di dalam abstrak."
    )
    set_font(run_abs_id, size_pt=10)

    p_key_id = doc.add_paragraph()
    p_key_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_key_id.paragraph_format.left_indent = Inches(0.5)
    p_key_id.paragraph_format.right_indent = Inches(0.5)
    p_key_id.paragraph_format.space_after = Pt(18)
    run_key_lbl_id = p_key_id.add_run("Kata Kunci: ")
    set_font(run_key_lbl_id, size_pt=9, bold=True)
    run_key_val_id = p_key_id.add_run("Arial 9 pt, template naskah, jurnal ilmiah, publikasi cepat (3-5 kata kunci, dipisahkan koma)")
    set_font(run_key_val_id, size_pt=9, italic=True)

    # 4. Abstract Section (English)
    p_abs_title_en = doc.add_paragraph()
    p_abs_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs_title_en.paragraph_format.space_before = Pt(6)
    p_abs_title_en.paragraph_format.space_after = Pt(6)
    run_abs_title_en = p_abs_title_en.add_run("ABSTRACT")
    set_font(run_abs_title_en, size_pt=10, bold=True, color_rgb=RGBColor(0x1A, 0x20, 0x2C))

    p_abs_en = doc.add_paragraph()
    p_abs_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs_en.paragraph_format.left_indent = Inches(0.5)
    p_abs_en.paragraph_format.right_indent = Inches(0.5)
    p_abs_en.paragraph_format.space_after = Pt(6)
    p_abs_en.paragraph_format.line_spacing = 1.0
    run_abs_en = p_abs_en.add_run(
        "The abstract must be written in English. It should be structured in a single paragraph, ranging from 150 to 250 words. "
        "It should briefly and accurately state the purpose of the research, the methods used, the main findings, and the principal conclusions. "
        "Use Arial 10 pt with single line spacing. Ensure clear grammar, scientific terminology, and avoid citations or mathematical formulas."
    )
    set_font(run_abs_en, size_pt=10, italic=True)

    p_key_en = doc.add_paragraph()
    p_key_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_key_en.paragraph_format.left_indent = Inches(0.5)
    p_key_en.paragraph_format.right_indent = Inches(0.5)
    p_key_en.paragraph_format.space_after = Pt(24)
    run_key_lbl_en = p_key_en.add_run("Keywords: ")
    set_font(run_key_lbl_en, size_pt=9, bold=True)
    run_key_val_en = p_key_en.add_run("Arial 9 pt, manuscript template, scientific journal, rapid publication (3-5 keywords, comma-separated)")
    set_font(run_key_val_en, size_pt=9, italic=True)

    # 5. Main Body Headers and Paragraphs
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, size_pt=11, bold=True, color_rgb=RGBColor(0x1A, 0x20, 0x2C))
        return p

    def add_body_paragraph(text, first_line_indent=0.3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Inches(first_line_indent)
        run = p.add_run(text)
        set_font(run, size_pt=11)
        return p

    # PENDAHULUAN
    add_section_header("PENDAHULUAN")
    add_body_paragraph(
        "Pendahuluan memaparkan latar belakang penelitian secara ringkas, padat, dan langsung menuju ke inti permasalahan. "
        "Gunakan huruf Arial 11 pt dengan spasi 1,15. Paragraf ditulis menjorok ke dalam pada baris pertama sebesar 0,76 cm (0,3 inci). "
        "Bagian pendahuluan harus memuat perbandingan dengan penelitian terdahulu secara kritis untuk menunjukkan orisinalitas "
        "dan kebaruan ilmiah (novelty) dari penelitian Anda."
    )
    add_body_paragraph(
        "Selanjutnya, uraikan dengan tegas rumusan masalah atau hipotesis penelitian (jika ada), disusul dengan tujuan penelitian "
        "yang ingin dicapai di akhir pendahuluan. Penulisan rujukan pustaka di dalam teks sangat disarankan menggunakan manajer "
        "referensi otomatis seperti Mendeley atau Zotero dengan format kutipan APA (American Psychological Association), "
        "contoh: (Pratama & Haryono, 2025)."
    )

    # METODE
    add_section_header("METODE")
    add_body_paragraph(
        "Metode penelitian ditulis secara terperinci sehingga pembaca atau peneliti lain dapat mengulangi penelitian Anda secara "
        "akurat. Uraikan jenis dan rancangan penelitian, subjek atau populasi dan teknik penarikan sampel, instrumen yang digunakan "
        "(serta uji validitas dan reliabilitasnya jika diperlukan), metode pengumpulan data, serta teknik analisis data statistik "
        "atau deskriptif kualitatif yang diterapkan."
    )
    add_body_paragraph(
        "Gunakan rumus matematika secara selektif dan ketik menggunakan fitur Equation Editor resmi Microsoft Word, bukan disisipkan "
        "sebagai file gambar. Setiap variabel dalam rumus harus didefinisikan dengan jelas dalam teks paragraf."
    )

    # HASIL DAN PEMBAHASAN
    add_section_header("HASIL DAN PEMBAHASAN")
    add_body_paragraph(
        "Hasil penelitian menyajikan data ilmiah secara lugas dan sistematis. Pembahasan bukan sekadar menuliskan ulang data "
        "yang ada pada tabel atau gambar, melainkan menjelaskan interpretasi ilmiah atas temuan tersebut, membandingkannya "
        "dengan teori atau hasil penelitian terdahulu yang relevan, serta mengemukakan keterbatasan penelitian ini."
    )
    add_body_paragraph(
        "Tabel dan Gambar harus diletakkan sedekat mungkin dengan penyebutannya di dalam teks. Setiap Tabel dan Gambar harus diberi "
        "nomor urut dan judul yang jelas, serta dirujuk di dalam kalimat (contoh: '...sebagaimana disajikan pada Tabel 1')."
    )

    # Contoh Tabel
    p_tbl_caption = doc.add_paragraph()
    p_tbl_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_tbl_caption.paragraph_format.space_before = Pt(8)
    p_tbl_caption.paragraph_format.space_after = Pt(4)
    run_tbl_caption = p_tbl_caption.add_run("Tabel 1. Contoh Format Tabel Bersih Tanpa Garis Vertikal")
    set_font(run_tbl_caption, size_pt=9.5, bold=True)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["Kategori Variabel", "Frekuensi (N)", "Persentase (%)"]
    for i, title in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = title
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size_pt=9.5, bold=True)
            
    rows_data = [
        ["Kategori A (Sangat Baik)", "45", "45,0%"],
        ["Kategori B (Cukup Baik)", "35", "35,0%"],
        ["Kategori C (Kurang Baik)", "20", "20,0%"]
    ]
    for row_idx, row_val in enumerate(rows_data):
        for col_idx, text in enumerate(row_val):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = text
            p = cell.paragraphs[0]
            if col_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_font(r, size_pt=9.5)

    # Space after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(6)
    p_space.paragraph_format.space_after = Pt(6)

    # KESIMPULAN
    add_section_header("KESIMPULAN")
    add_body_paragraph(
        "Kesimpulan menyajikan rangkuman jawaban atas pertanyaan penelitian atau tujuan penulisan artikel. Kesimpulan disajikan "
        "dalam bentuk paragraf mengalir (bukan dalam bentuk poin-poin penomoran) dan ditulis secara ringkas tanpa mengulang data "
        "statistik secara berlebihan. Tambahkan pula saran atau rekomendasi praktis untuk penelitian selanjutnya di bagian akhir paragraf."
    )

    # DAFTAR PUSTAKA
    p_ref_header = doc.add_paragraph()
    p_ref_header.paragraph_format.space_before = Pt(18)
    p_ref_header.paragraph_format.space_after = Pt(8)
    p_ref_header.paragraph_format.keep_with_next = True
    run_ref = p_ref_header.add_run("DAFTAR PUSTAKA")
    set_font(run_ref, size_pt=11, bold=True, color_rgb=RGBColor(0x1A, 0x20, 0x2C))

    def add_reference_item(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3) # Hanging indent
        run = p.add_run(text)
        set_font(run, size_pt=9.5)

    add_reference_item(
        "Pratama, A. B., & Haryono, S. (2025). Analisis Efisiensi Tata Kelola Kebijakan Publik di Indonesia. "
        "Jurnal Audit Kebijakan Publik, 1(1), 12-25. https://doi.org/10.47822/rjrakp.jakp.v1i1.01"
    )
    add_reference_item(
        "Ramadhan, F., & Sitorus, T. (2024). Implementasi Algoritma Machine Learning untuk Klasifikasi Dokumen Jurnal. "
        "Jurnal Teknik dan Teknologi, 3(2), 88-97."
    )
    add_reference_item(
        "Santoso, B. (2023). Metode Penelitian Kuantitatif untuk Pemula (Edisi ke-2). Jakarta: Penerbit Akademika."
    )

    # Save to public directory
    out_dir = "public"
    if not os.path.exists(out_dir):
         os.makedirs(out_dir)
         
    output_path = os.path.join(out_dir, "template-naskah.docx")
    doc.save(output_path)
    print(f"SUCCESS: {output_path}")

if __name__ == "__main__":
    create_template()
